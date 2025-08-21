import json
import os

import docx
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
import pandas as pd
from docxtpl import DocxTemplate
# Мои импорты
import common_generator
import editors as ed

data: json

# S - Генератор Спецификации
cat_names_plural = ed.cat_names_plural
cat_names_singular = ed.cat_names_singular

dfs = []  # DataFrame'ы с изначальными данными
dict_chars = {}  # Словарь комбинированных изначальных DataFrame'ов с их буквенными обозначениями
final_df = []  # Список финальных DataFrame'ов на вставку в документ(-ы)
names_df = {}  # Словарь с формированными строками для вставки в документы

files = []  # Массив названий файлов для обработки
input_files_folder: str  # Путь к файлам

prim_cats = {}  # Категории примечаний
prim_not_install = {}  # Категории для "Не устанавливать"

components_one_manuf_categories = {}
civ: bool

def split_to_format(string: str, shift_threshold):
    i = 0  # Индекс для добавления части наименования в уже существующую строку
    paste_string = ['']
    for name_part in string.split():
        # Если длина итоговой строки длиннее порога, то создаем новый перенос
        if len(f"{paste_string[-1]} {name_part}") > shift_threshold:
            paste_string.append(name_part)
            i += 1
            continue
        paste_string[i] += f" {name_part}"

    for index, string in enumerate(paste_string):
        paste_string[index] = paste_string[index].lstrip().rstrip()

    return paste_string

def export_to_word():
    """
    Экспортирует полученный список в формат Word
    """
    if data['Templates_Path'] != "":
        path_to_template = data['Templates_Path']
    else:
        path_to_template = common_generator.prog_dir + "\\Шаблоны"

    if not civ:
        doc = Document(path_to_template + '\\Шаблон СП.docx')
    else:
        doc = Document(path_to_template + '\\Шаблон СП Гражданский.docx')

    style = doc.styles['Normal']
    font = style.font
    font.name = 'T-FLEX Type A'
    font.size = Pt(12)

    module = dict_chars['C'][0].module

    # Получаем итераторы таблиц в документе и строк в нем
    tables = iter(doc.tables)
    table = next(tables)

    rows = iter(table.rows)
    start_row = True

    try:
        # Заполнение первой страницы прилагаемой документацией
        next(rows)
        next(rows)
        row = next(rows)

        row_index = 2

        row.cells[4].text = "Документация"
        row.cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row.cells[4].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE

        next(rows)
        row = next(rows)

        row_index += 2

        row.cells[0].text = "А1"
        row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        row.cells[4].text = "Сборочный чертеж"
        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        paste_string = split_to_format((module + " СБ"), 34)

        for string in paste_string:
            row.cells[3].text = string
            row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            row = next(rows)
            row_index += 1

        if data['Scheme']:
            row.cells[0].text = "А3"
            row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            row.cells[4].text = "Схема электрическая"
            row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            paste_string = split_to_format((module + " Э3"), 34)

            for index, string in enumerate(paste_string):
                row.cells[3].text = string
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

                if index == 1:
                    row.cells[4].text = "принципиальная"
                    row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            if index == 0:
                row.cells[4].text = "принципиальная"
                row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

        if data['PE']:
            row.cells[0].text = "А4"
            row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            row.cells[4].text = "Перечень элементов"
            row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            paste_string = split_to_format((module + " ПЭ3"), 34)
            for string in paste_string:
                row.cells[3].text = string
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

        if data['DK']:
            row.cells[0].text = "А4"
            row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            row.cells[4].text = "Комплект карт для оценки"
            row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            paste_string = split_to_format((module + " ДК"), 34)
            for index, string in enumerate(paste_string):
                row.cells[3].text = string
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

                if index == 1:
                    row.cells[4].text = "правильности применения ЭРИ"
                    row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            if index == 0:
                row.cells[4].text = "правильности применения ЭРИ"
                row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

        if data['I1']:
            row.cells[0].text = "А4"
            row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            row.cells[4].text = "Инструкция по"
            row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            paste_string = split_to_format((module + " И1"), 34)
            for index, string in enumerate(paste_string):
                row.cells[3].text = string
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

                if index == 1:
                    row.cells[4].text = "программированию"
                    row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            if index == 0:
                row.cells[4].text = "программированию"
                row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

        if data['I2']:
            row.cells[0].text = "А4"
            row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            row.cells[4].text = "Инструкция по настройке"
            row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

            paste_string = split_to_format((module + " И2"), 34)
            for string in paste_string:
                row.cells[3].text = string
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                row = next(rows)
                row_index += 1

        next(rows)
        row = next(rows)
        row_index += 2

        row.cells[4].text = "Сборочные единицы"
        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row.cells[4].paragraphs[0].runs[0].underline = \
            docx.enum.text.WD_UNDERLINE.SINGLE

        next(rows)
        row = next(rows)
        row_index += 2

        row.cells[3].text = module
        row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        row.cells[2].text = '1'
        row.cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        row.cells[4].text = "Плата печатная"
        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        row.cells[5].text = '1'
        row.cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        next(rows)
        next(rows)
        row = next(rows)
        row_index += 3

        row.cells[4].text = "Прочие изделия"
        row.cells[4].paragraphs[0].runs[0].underline = \
            docx.enum.text.WD_UNDERLINE.SINGLE

        chip_pos = 1
        for char in sorted(dict_chars.keys()):
            for chip_index, chip in enumerate(dict_chars[char]):
                chip_pos += 1
                chip.split_name(shift_threshold=32)
                chip.split_designators(shift_threshold=59)

                if chip_index == 0:
                    try:
                        row = next(rows)
                        row_index += 1
                    except StopIteration:
                        table = next(tables)
                        rows = iter(table.rows)
                        next(rows)
                        row = next(rows)

                        row_index = 1

                    rows_len = len(table.rows) - 1
                    if row_index + len(chip.name) > rows_len or row_index + len(chip.designator) > rows_len \
                            or (isinstance(chip.quantity, list) and row_index + len(chip.quantity) > rows_len) or \
                            (isinstance(chip.manuf, list) and row_index + len(chip.manuf) > rows_len):
                        table = next(tables)
                        rows = iter(table.rows)
                        next(rows)
                        row = next(rows)

                        row_index = 1

                    cat_name = ''
                    if len(dict_chars[char]) > 1:
                        for designator, d_cat_name in cat_names_plural.items():
                            if char == designator:
                                cat_name = d_cat_name

                                if char in components_one_manuf_categories.keys():
                                    cat_name += f', {components_one_manuf_categories[char]}'
                                break

                        row.cells[4].text = cat_name
                        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        row.cells[4].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE

                        try:
                            row = next(rows)
                            row_index += 1
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            row = next(rows)

                            row_index = 1
                    # Если компонент один
                    else:
                        for designator, d_cat_name in cat_names_singular.items():
                            if char == designator:
                                cat_name = d_cat_name

                        chip.split_name(shift_threshold=32, cat_name=cat_name)

                # Вставка строки в документ и ее оформление
                rows_len = len(table.rows) - 1
                if row_index + len(chip.name) - 1 > rows_len or row_index + len(chip.designator) - 1 > rows_len \
                        or (isinstance(chip.quantity, list) and row_index + len(chip.quantity) - 1 > rows_len) or \
                        (isinstance(chip.manuf, list) and row_index + len(chip.manuf) - 1 > rows_len):
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    row = next(rows)

                    row_index = 1

                row.cells[2].text = str(chip_pos)
                row.cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                row.cells[5].text = str(chip.quantity)
                row.cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                while True:
                    try:
                        row.cells[6].text = str(chip.designator[0])
                        row.cells[6].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                        del chip.designator[0]
                    except IndexError:
                        pass

                    try:
                        row.cells[4].text = str(chip.name[0])
                        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                        del chip.name[0]
                    except IndexError:
                        pass

                    if chip.name != [] or chip.designator != []:
                        try:
                            row = next(rows)
                            row_index += 1
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            row = next(rows)

                            row_index = 1
                    else:
                        break

                try:
                    row = next(rows)
                    row_index += 1
                except StopIteration:
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    row = next(rows)

                    row_index = 1
    except StopIteration:
        pass

    # Заполнение Примечаний

    # Сохранение документа при заполнении таблицы
    save_dir = files[0][:-len(files[0].split("\\")[-1])] + "\\Спецификации\\"

    new_name = pd.read_excel(files[0], sheet_name='BOM', header=None).loc[7, 10].split(' ')[0]
    if new_name == '':
        print("Не заполнено поле 'Первичная Применяемость'!")

    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    doc.save(save_dir + new_name + ' СП (без полей).docx')

    doc = DocxTemplate(save_dir + new_name + ' СП (без полей).docx')

    device_name = pd.read_excel(files[0], sheet_name='BOM', header=None).loc[2, 0]

    context = {"DocName": new_name + ' СП', "PervPrim": new_name, "Razrab": data['Razrab'],
               "Proveril": data['Proveril'], "N_control": data["N_control"], "Utverdil": data['Utverdil'],
               "DeviceName": device_name, "PlateName": new_name.split(' ')[0]}
    doc.render(context)

    if civ:
        new_name += ' гражданская'
    doc.save(save_dir + new_name + ' СП.docx')
    print(f"Файл сохранен по пути: {save_dir + new_name + ' СП.docx'}")
    return new_name.split(' ')[0]


def execute(input_files: list, is_civ: bool):
    global dfs, files, dict_chars, prim_not_install, prim_cats, final_df, data, civ, components_one_manuf_categories
    civ = is_civ
    files = input_files
    with open("Profile.json", "r") as read_file:
        data = json.load(read_file)

    i = 1
    for file in files:
        print(f"\nФормирую {i}-ю СП\n"
              f"=============")
        i += 1
        dfs = []
        dict_chars = {}

        common_generator.test = True

        common_generator.no_perechen = 1
        print("Получаю данные из перечней элементов...")
        dfs, files = common_generator.get_dfs(dict_chars, [file])
        print("Формирую общую таблицу элементов...")
        dict_chars, prim_not_install, prim_cats, one_man_cats = common_generator.get_components(dict_chars, dfs, [file])
        print("Проверяю примечания на регулирование...")
        dict_chars = common_generator.split_to_regul(dict_chars)
        print("Комбинирую компоненты...")
        dict_chars = common_generator.combine_chips_in_module(dict_chars)
        print("Комбинирую модули...")
        dict_chars = common_generator.combine_modules(dict_chars)
        print("Вставляю готовый перечень в шаблон...")
        export_to_word()
    print("\n=============\n"
          "Готово!\n\n\n")
