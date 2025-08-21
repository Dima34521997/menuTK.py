import json
import os

import docx
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
import pandas as pd
from docxtpl import DocxTemplate

# Мои импорты
import obshiy_perechen
import editors as ed

# Генератор ПЭ
cat_names_plural = ed.cat_names_plural
cat_names_singular = ed.cat_names_singular

dfs = []  # DataFrame'ы с изначальными данными
dict_chars = {}  # Словарь комбинированных изначальных DataFrame'ов с их буквенными обозначениями
final_df = []  # Список финальных DataFrame'ов на вставку в документ(-ы)
names_df = {}  # Словарь с формированными строками для вставки в документы

files = []  # Массив названий файлов для обработки
input_files_folder: str  # Путь к файлам

prim_cats = {}  # Категории примечаний
prim_not_install = {}  # Категории для "Не устанавливать"

components_one_manuf_categories = {}

data: json
civ: bool


def export_to_word():
    """
    Экспортирует полученный список в формат Word
    """
    if data['Templates_Path'] != "":
        path_to_template = data['Templates_Path']
    else:
        path_to_template = obshiy_perechen.prog_dir + "\\Шаблоны"

    if not civ:
        doc = Document(path_to_template + '\\Шаблон ПЭ.docx')
    else:
        doc = Document(path_to_template + '\\Шаблон ПЭ Гражданский.docx')

    style = doc.styles['Normal']
    font = style.font
    font.name = 'T-FLEX Type A'
    font.size = Pt(12)

    # Получаем итератори таблиц в документе и строк в нем
    tables = iter(doc.tables)
    table = next(tables)

    rows = iter(table.rows)
    next(rows)
    row = next(rows)
    row_index = 1
    try:
        for char in sorted(dict_chars.keys()):
            for component_index, component in enumerate(dict_chars[char]):
                component.split_name(shift_threshold=52)
                component.split_designators(shift_threshold=53)
                component.split_notice(shift_threshold=120)

                # Вставка наименования категории компонентов
                if component_index == 0:
                    try:
                        row = next(rows)
                        row_index += 1
                    except StopIteration:
                        table = next(tables)
                        rows = iter(table.rows)
                        next(rows)
                        row = next(rows)

                        row_index = 1

                    rows_len = len(table.rows) - 1
                    if row_index + len(component.name) > rows_len or row_index + len(component.designator) > rows_len \
                            or (isinstance(component.quantity, list) and row_index + len(component.quantity) > rows_len) or \
                            (isinstance(component.manuf, list) and row_index + len(component.manuf) > rows_len):
                        table = next(tables)
                        rows = iter(table.rows)
                        next(rows)
                        row = next(rows)

                        row_index = 1

                    cat_name = ''
                    if len(dict_chars[char]) > 1:
                        for desig, d_cat_name in cat_names_plural.items():
                            if char == desig:
                                cat_name = d_cat_name

                                if char in components_one_manuf_categories.keys():
                                    cat_name += f', {components_one_manuf_categories[char]}'
                                break

                        row.cells[1].text = cat_name
                        row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        row.cells[1].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE

                        try:
                            row = next(rows)
                            row_index += 1
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            row = next(rows)

                            row_index = 1
                    # Если компонент один
                    else:
                        for desig, d_cat_name in cat_names_singular.items():
                            if char == desig:
                                cat_name = d_cat_name

                        component.split_name(shift_threshold=56, cat_name=cat_name)

                # Вставка строки в документ и ее оформление
                rows_len = len(table.rows) - 1
                if row_index + len(component.name) - 1 > rows_len or row_index + len(component.designator) - 1 > rows_len \
                        or (
                        isinstance(component.quantity, list) and row_index + len(component.quantity) - 1 > rows_len) or \
                        (isinstance(component.manuf, list) and row_index + len(component.manuf) - 1 > rows_len):
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    row = next(rows)

                    row_index = 1

                row.cells[2].text = str(component.quantity)
                row.cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                row.cells[3].text = component.notice
                row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                # Вставка примечаний построчно
                while True:
                    # Вставляем десиг, наименование, кол-во и примечание по одной строке
                    try:
                        row.cells[0].text = str(component.designator[0])
                        row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        del component.designator[0]
                    except IndexError:
                        pass

                    try:
                        row.cells[1].text = str(component.name[0])
                        row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                        del component.name[0]
                    except IndexError:
                        pass

                    # --- Вставка примечания построчно ---
                    try:
                        row.cells[3].text = str(component.notice[0])
                        row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                        del component.notice[0]
                    except IndexError:
                        row.cells[3].text = ""
                    # ---

                    # Для компонентов R и C — вставка man_part_num отдельно
                    if (char == 'C' or char == 'R') and component.manuf.find("ТУ", 0) == -1:
                        try:
                            row.cells[3].text = str(component.man_part_num)
                            row.cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                            del component.name[0]  # Удаляем запятую после man_part_num
                        except IndexError:
                            pass

                    # Переход к следующей строке, если остались данные
                    if component.name != [] or component.designator != [] or component.notice != []:
                        try:
                            row = next(rows)
                            row_index += 1
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            row = next(rows)
                            row_index = 1
                    else:
                        break
                try:
                    row = next(rows)
                    row_index += 1
                except StopIteration:
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    row = next(rows)
                    start_row = True

                    row_index = 1

    except StopIteration:
        pass

    save_dir = files[0][:-len(files[0].split("\\")[-1])] + "\\Перечни\\"

    new_name = pd.read_excel(files[0], sheet_name='BOM', header=None).loc[7, 10].split(' ')[0]
    deviceName = pd.read_excel(files[0], sheet_name='BOM', header=None).loc[2, 0]
    if new_name == '':
        print("Не заполнено поле 'Первичная Применямость'!")
    if deviceName == '':
        print("Не заполнено поле 'Наименование1'!")

    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    doc.save(save_dir + new_name + ' ПЭ (без полей).docx')

    doc = DocxTemplate(save_dir + new_name + ' ПЭ (без полей).docx')

    context = {"DocName": new_name + ' ПЭ', "PervPrim": new_name, "Razrab": data['Razrab'],
               "Proveril": data['Proveril'], "N_control": data["N_control"], "Utverdil": data['Utverdil'],
               "DeviceName": deviceName, "PlateName": new_name.split(' ')[0]}
    doc.render(context)

    if civ:
        new_name += ' гражданский'
    doc.save(save_dir + new_name + ' ПЭ.docx')
    print(f"Файл сохранен по пути: {save_dir + new_name + '.docx'}")


def execute(input_files: list, is_civ: bool):
    global dfs, files, dict_chars, prim_not_install, prim_cats, final_df, data, civ, components_one_manuf_categories
    civ = is_civ
    files = input_files
    with open("Profile.json", "r") as read_file:
        data = json.load(read_file)

    i = 1
    for file in files:
        print(f"\nФормирую {i}-й ПЭ\n"
              f"=============")
        i += 1
        dfs = []
        dict_chars = {}

        obshiy_perechen.test = True

        obshiy_perechen.no_perechen = 0
        print("Получаю данные из перечней элементов...")
        dfs, files = obshiy_perechen.get_dfs(dict_chars, [file])
        print("Формирую общую таблицу элементов...")
        dict_chars, prim_not_install, prim_cats, \
            components_one_manuf_categories = obshiy_perechen.get_components(dict_chars, dfs, [file])
        print("Проверяю примечения на регулирование...")
        dict_chars = obshiy_perechen.split_to_regul(dict_chars)
        print("Вставляю готовый перечень в шаблон...")
        export_to_word()
    print("\n=============\n"
          "Готово!\n\n\n")
