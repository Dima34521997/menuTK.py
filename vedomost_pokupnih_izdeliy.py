import json
import os

import docx
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import pandas as pd

# Мои импорты
from docxtpl import DocxTemplate

import obshiy_perechen
import editors as ed

# Программа для создания ВП

cat_names_plural = ed.cat_names_plural
cat_names_singular = ed.cat_names_singular

dfs = []  # DataFrame'ы с изначальными данными
dict_chars = {}  # Словарь комбинированных изначальных DataFrame'ов с их буквенными обозначениями
final_df = []  # Список финальных DataFrame'ов на вставку в документ(-ы)
names_df = {}  # Словарь с формированными строками для вставки в документы

files = []  # Массив названий файлов для обработки
input_files_folder: str  # Путь к файлам

prim_cats = {}  # Категории примечаний
prim_not_install = {}  # Категории для "Не устанавливать"

components_one_manuf_categories = {}

data: json


def export_to_word():
    """
    Экспортирует полученный список в формат Word
    """

    if data['Templates_Path'] != "":
        path_to_template = data['Templates_Path']
    else:
        path_to_template = obshiy_perechen.prog_dir + "\\Шаблоны"

    doc = Document(path_to_template + '\\Шаблон ВП.docx')

    style = doc.styles['Normal']
    font = style.font
    font.name = 'T-FLEX Type A'
    font.size = Pt(12)

    # Получаем итератори таблиц в документе и строк в нем
    tables = iter(doc.tables)
    next(tables)
    table = next(tables)

    rows = iter(table.rows)
    next(rows)
    row = next(rows)
    start_row = True

    try:
        for char in sorted(dict_chars.keys()):
            for chip_index, chip in enumerate(dict_chars[char]):
                man = chip.manuf
                chip.manuf = ''
                # Перезаписываю наименование уже без производителя
                chip.name = chip._make_name()[:-2]

                chip.split_name(shift_threshold=32)
                chip.designator = []

                # Вставка наименования категории компонентов
                if chip_index == 0:
                    rows_len = len(table.rows) - 2
                    # Если начало таблицы
                    if start_row:
                        row = next(rows)
                        start_row = False

                        row_index = int(row.cells[0].paragraphs[0].runs[0].text)
                        if row_index + len(chip.name) - 1 > rows_len or row_index + len(chip.designator) - 1 > rows_len \
                                or (
                                isinstance(chip.quantity, list) and row_index + len(chip.quantity) - 1 > rows_len) or \
                                (isinstance(chip.manuf, list) and row_index + len(chip.manuf) - 1 > rows_len):
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            next(rows)
                            # Начинаем каждый раз со второй строки
                            row = next(rows)

                            start_row = True
                    # Если не начало таблицы
                    else:
                        # Начало категории не должно стоять в конце таблицы,
                        # поэтому прибавляю еще одну строку в учет индекса i
                        try:
                            row = next(rows)
                            start_row = False

                            row_index = int(row.cells[0].paragraphs[0].runs[0].text)
                            if row_index + len(chip.name) - 1 > rows_len or row_index + len(chip.designator) - 1 > rows_len \
                                    or (isinstance(chip.quantity, list) and row_index + len(
                                chip.quantity) - 1 > rows_len) or \
                                    (isinstance(chip.manuf, list) and row_index + len(chip.manuf) - 1 > rows_len):
                                table = next(tables)
                                rows = iter(table.rows)
                                next(rows)
                                next(rows)
                                # Начинаем каждый раз со второй строки
                                row = next(rows)

                                start_row = True
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            next(rows)
                            # Начинаем каждый раз со второй строки
                            row = next(rows)

                            start_row = True

                    # Формирование наименования категории
                    cat_name = ''
                    # Если компонент в категории не один
                    if len(dict_chars[char]) > 1:
                        for desig, d_cat_name in cat_names_plural.items():
                            if char == desig:
                                cat_name = d_cat_name
                                break

                        row.cells[1].text = cat_name
                        row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        row.cells[1].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE

                        try:
                            row = next(rows)
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            next(rows)
                            # Начинаем каждый раз со второй строки
                            row = next(rows)
                    # Если все-таки один
                    else:
                        for desig, d_cat_name in cat_names_singular.items():
                            if char == desig:
                                cat_name = d_cat_name

                        # Наименование категории и компонента в одной строке
                        chip.split_name(shift_threshold=33, cat_name=cat_name)

                # Вставка строки в документ и ее оформление
                # Если количество строк не хватает (для переносов или модулей), то переходим на следующую таблицу
                chip.manuf = man
                chip.split_man(shift_threshold=33)

                rows_len = len(table.rows) - 2
                row_index = int(row.cells[0].paragraphs[0].runs[0].text)
                if row_index + len(chip.name) - 1 > rows_len or row_index + len(chip.designator) - 1 > rows_len \
                        or (isinstance(chip.quantity, list) and row_index + len(chip.quantity) - 1 > rows_len) or \
                        (isinstance(chip.manuf, list) and row_index + len(chip.manuf) - 1 > rows_len):
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    next(rows)
                    # Начинаем каждый раз со второй строки
                    row = next(rows)

                if chip.char == 'R' or chip.char == 'C':
                    row.cells[2].text = str(chip.man_part_num)
                    row.cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                row.cells[6].text = str(chip.quantity)
                row.cells[6].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                row.cells[9].text = str(chip.quantity)
                row.cells[9].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                row.cells[10].text = chip.notice
                row.cells[10].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                while True:
                    try:
                        row.cells[1].text = str(chip.name[0])
                        row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                        del chip.name[0]
                    except IndexError:
                        pass

                    try:
                        row.cells[4].text = str(chip.manuf[0])
                        row.cells[4].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        del chip.manuf[0]
                    except IndexError:
                        pass
                    except TypeError:
                        row.cells[5].text = str(chip.manuf)
                        row.cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        chip.manuf = ''

                    try:
                        row.cells[5].text = str(chip.module[0])
                        row.cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        del chip.module[0]
                    except IndexError:
                        pass
                    except TypeError:
                        row.cells[5].text = str(chip.module)
                        row.cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        chip.module = ''

                    try:
                        row.cells[6].text = str(chip.quantity[0])
                        row.cells[6].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        row.cells[9].text = str(chip.quantity[0])
                        row.cells[9].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        if len(chip.quantity) == 1:
                            row.cells[6].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE
                            row.cells[9].paragraphs[0].runs[0].underline = docx.enum.text.WD_UNDERLINE.SINGLE
                        del chip.quantity[0]
                    except IndexError:
                        pass
                    except TypeError:
                        row.cells[5].text = str(chip.quantity)
                        row.cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        chip.quantity = ''

                    if chip.name != [] or chip.designator != [] or \
                            (isinstance(chip.quantity, list) and chip.quantity != []) or \
                            (isinstance(chip.module, list) and chip.module != []) or \
                            (isinstance(chip.manuf, list) and chip.manuf != []):
                        try:
                            row = next(rows)
                        except StopIteration:
                            table = next(tables)
                            rows = iter(table.rows)
                            next(rows)
                            next(rows)
                            # Начинаем каждый раз со второй строки
                            row = next(rows)
                    else:
                        break

                try:
                    row = next(rows)
                except StopIteration:
                    table = next(tables)
                    rows = iter(table.rows)
                    next(rows)
                    next(rows)
                    # Начинаем каждый раз со второй строки
                    row = next(rows)
    except StopIteration:
        pass

    # Если один BOM, то будет использоваться код этого изделия
    # Еслин несколько, то ввод вручную
    if len(files) > 1:
        new_name = data['Project_Name']
    else:
        new_name = pd.read_excel(files[0], sheet_name='BOM', header=None).loc[7, 10].split(' ')[0]
        if new_name == '':
            print("Не заполнено поле 'Первичная Применямость'!")

    save_dir = files[0][:-len(files[0].split("\\")[-1])] + "\\ВП\\"

    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    doc.save(save_dir + new_name + ' ВП.docx')

    doc = DocxTemplate(save_dir + new_name + ' ВП.docx')

    # Вставка значений в поля в форме Word
    context = {"PlateName": new_name, "PervPrim": new_name, "Razrab": data['Razrab'],
               "Proveril": data['Proveril'], "N_control": data["N_control"], "Utverdil": data['Utverdil']}
    doc.render(context)
    doc.save(save_dir + new_name + ' ВП.docx')
    print(f"Файл сохранен по пути: {save_dir + new_name + ' ВП.docx'}")


def execute(input_files):
    global dfs, files, dict_chars, prim_not_install, prim_cats, final_df, data, components_one_manuf_categories
    files = input_files
    with open("Profile.json", "r") as read_file:
        data = json.load(read_file)

    dfs = []
    dict_chars = {}

    obshiy_perechen.test = True

    obshiy_perechen.no_perechen = 1
    print("Получаю данные из перечней элементов...")
    dfs, files = obshiy_perechen.get_dfs(dict_chars, files)
    print("Формирую общую таблицу элементов...")
    dict_chars, prim_not_install, prim_cats, one_man_cats = obshiy_perechen.get_components(dict_chars, dfs, files)
    print("Проверяю примечения на регулирование...")
    dict_chars = obshiy_perechen.split_to_regul(dict_chars)
    print("Комбинирую компоненты...")
    dict_chars = obshiy_perechen.combine_chips_in_module(dict_chars)
    print("Комбинирую модули...")
    dict_chars = obshiy_perechen.combine_modules(dict_chars)
    print("Вставляю готовый перечень в шаблон...")
    export_to_word()

    print("\n=============\n"
          "Готово!\n\n\n")
